"""Helpers to build in-memory .docx fixtures for tests.

Tests construct documents from a list of paragraph strings (optionally tagged
with a Word style name) so no binary fixture files need to be committed for the
common cases.
"""

from __future__ import annotations

import io

from docx import Document


def make_docx(
    paragraphs: list[str | tuple[str, str]],
    tables: list[list[list[str]]] | None = None,
) -> bytes:
    """Build a .docx from paragraph specs and return its bytes.

    Each paragraph item is either a plain string (Normal style) or a
    ``(style, text)`` tuple, where ``style`` is a Word style name such as
    ``"Heading 1"``. ``tables`` is an optional list of tables, each a list of
    rows, each row a list of cell strings; they are appended after the
    paragraphs.
    """
    document = Document()
    for para in paragraphs:
        if isinstance(para, tuple):
            style, text = para
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(para)
    for rows in tables or []:
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                table.rows[r].cells[c].text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
